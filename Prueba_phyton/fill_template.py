#!/usr/bin/env python3
"""
fill_template.py
================
Rellena automáticamente un template .docx con datos de usuario e imágenes.

PLACEHOLDERS en el .docx
─────────────────────────
  {{nombre}}      → texto  (cualquier clave del frontmatter YAML)
  {{IMAGEN_1}}    → imagen (nombre de sección de imágenes en el .md)

FORMATO del archivo .md de datos
─────────────────────────────────
  ---
  nombre: Juan Pérez
  cargo: Técnico de Soporte
  equipo: PC-LAB-03
  fecha: 2026-05-14
  sistema_operativo: Windows 11
  ip: 192.168.1.50
  ---

  imagenes:
    IMAGEN_1: capturas/escritorio.png
    IMAGEN_2: capturas/sistema.png
    IMAGEN_3: /ruta/absoluta/red.png

  Las rutas de imagen pueden ser absolutas o relativas al directorio del .md.

USO
───
  # Básico (genera salida junto al .md)
  python fill_template.py -t formato.docx -d usuario.md

  # Con salida específica
  python fill_template.py -t formato.docx -d usuario.md -o resultado/juan.docx

  # Ajustar ancho máximo de imágenes (por defecto 15 cm)
  python fill_template.py -t formato.docx -d usuario.md --ancho 12
"""

import argparse
import re
import sys
from copy import deepcopy
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm


# ─────────────────────────────────────────────
#  PARSING DEL MARKDOWN
# ─────────────────────────────────────────────

def parse_markdown(md_path: Path) -> tuple[dict, dict]:
    """
    Lee el archivo .md y devuelve:
      - datos: dict con el frontmatter YAML (datos del usuario)
      - imagenes: dict {PLACEHOLDER -> Path de imagen}
    """
    text = md_path.read_text(encoding="utf-8")

    # Extraer frontmatter YAML
    fm_match = re.match(r"^---\s*\n(.*?)\n---", text, re.DOTALL)
    if not fm_match:
        sys.exit(
            "ERROR: El archivo .md debe tener frontmatter YAML entre --- al inicio.\n"
            "Ejemplo:\n---\nnombre: Juan\n---"
        )
    datos = yaml.safe_load(fm_match.group(1)) or {}

    # Extraer sección 'imagenes:'
    img_match = re.search(r"^imagenes:\s*\n((?:[ \t]+\S.*\n?)+)", text, re.MULTILINE)
    imagenes_raw = {}
    if img_match:
        imagenes_raw = yaml.safe_load("imagenes:\n" + img_match.group(0).replace("imagenes:", ""))
        if isinstance(imagenes_raw, dict):
            imagenes_raw = imagenes_raw.get("imagenes", {}) or {}

    # Resolver rutas de imágenes (relativas al directorio del .md)
    base_dir = md_path.parent
    imagenes = {}
    for key, val in (imagenes_raw or {}).items():
        p = Path(str(val))
        if not p.is_absolute():
            p = base_dir / p
        imagenes[str(key).strip()] = p

    return datos, imagenes


# ─────────────────────────────────────────────
#  HELPERS DE PÁRRAFO
# ─────────────────────────────────────────────

def get_full_text(paragraph) -> str:
    """Texto completo de un párrafo (une todos los runs)."""
    return "".join(r.text for r in paragraph.runs)


def set_full_text(paragraph, new_text: str):
    """
    Reemplaza el texto de un párrafo conservando el formato del primer run.
    Vacía runs adicionales.
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_paragraph_with_image(paragraph, image_path: Path, max_width_cm: float):
    """
    Elimina el contenido de texto de un párrafo y pone una imagen en su lugar.
    """
    # Eliminar todos los runs del XML del párrafo
    p_xml = paragraph._p
    for r in list(p_xml.findall(qn("w:r"))):
        p_xml.remove(r)
    # Insertar imagen en un run nuevo
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(max_width_cm))


# ─────────────────────────────────────────────
#  ITERACIÓN POR TODOS LOS PÁRRAFOS
# ─────────────────────────────────────────────

def iter_paragraphs(document):
    """
    Genera todos los párrafos del documento:
    cuerpo principal + celdas de tablas (anidadas).
    """
    # Párrafos del cuerpo principal
    yield from document.paragraphs

    # Párrafos dentro de tablas (incluyendo tablas anidadas)
    def _from_table(table):
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested_table in cell.tables:
                    yield from _from_table(nested_table)

    for table in document.tables:
        yield from _from_table(table)


# ─────────────────────────────────────────────
#  REEMPLAZOS
# ─────────────────────────────────────────────

PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")


def apply_text_replacements(document, datos: dict):
    """
    Reemplaza todos los {{clave}} por su valor de texto en el documento.
    Maneja el caso en que el placeholder quede dividido entre varios runs.
    """
    # Normalizar claves a minúsculas para comparación flexible
    datos_norm = {k.lower(): str(v) for k, v in datos.items()}
    reemplazos = 0

    for paragraph in iter_paragraphs(document):
        full = get_full_text(paragraph)
        if "{{" not in full:
            continue

        def _sub(m):
            nonlocal reemplazos
            key = m.group(1).lower()
            if key in datos_norm:
                reemplazos += 1
                return datos_norm[key]
            return m.group(0)  # dejar tal cual si no se encuentra

        new_text = PLACEHOLDER_RE.sub(_sub, full)
        if new_text != full:
            set_full_text(paragraph, new_text)

    return reemplazos


def apply_image_replacements(document, imagenes: dict, max_width_cm: float):
    """
    Reemplaza todos los {{CLAVE_IMAGEN}} por la imagen correspondiente.
    """
    reemplazos = 0
    errores = []

    for paragraph in iter_paragraphs(document):
        full = get_full_text(paragraph).strip()
        m = PLACEHOLDER_RE.fullmatch(full)
        if not m:
            # También detectar si el párrafo CONTIENE solo el placeholder
            m = re.fullmatch(r"\s*\{\{(\w+)\}\}\s*", full)
        if not m:
            continue

        key = m.group(1)
        if key not in imagenes:
            continue  # Puede ser un placeholder de texto, ya manejado

        img_path = imagenes[key]
        if not img_path.exists():
            errores.append(f"  ⚠  Imagen no encontrada para {{{{key}}}}: {img_path}")
            continue

        replace_paragraph_with_image(paragraph, img_path, max_width_cm)
        reemplazos += 1
        print(f"  ✓  {{{{  {key}  }}}} → {img_path.name}")

    return reemplazos, errores


# ─────────────────────────────────────────────
#  MAIN
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Rellena un template .docx con datos de usuario e imágenes.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("-t", "--template", required=True, metavar="TEMPLATE.docx",
                        help="Ruta al archivo .docx de plantilla")
    parser.add_argument("-d", "--datos", required=True, metavar="USUARIO.md",
                        help="Archivo .md con frontmatter YAML y sección de imágenes")
    parser.add_argument("-o", "--salida", metavar="SALIDA.docx",
                        help="Ruta del archivo de salida (por defecto: <datos>_relleno.docx)")
    parser.add_argument("--ancho", type=float, default=15.0, metavar="CM",
                        help="Ancho máximo de imágenes en cm (por defecto: 15)")
    args = parser.parse_args()

    template_path = Path(args.template)
    datos_path = Path(args.datos)

    if not template_path.exists():
        sys.exit(f"ERROR: No se encontró el template: {template_path}")
    if not datos_path.exists():
        sys.exit(f"ERROR: No se encontró el archivo de datos: {datos_path}")

    # Ruta de salida
    if args.salida:
        salida_path = Path(args.salida)
    else:
        salida_path = datos_path.with_name(datos_path.stem + "_relleno.docx")

    salida_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"\n📄 Template  : {template_path}")
    print(f"📋 Datos     : {datos_path}")
    print(f"💾 Salida    : {salida_path}")
    print()

    # Leer datos
    datos, imagenes = parse_markdown(datos_path)

    print(f"Datos del usuario ({len(datos)} campos):")
    for k, v in datos.items():
        print(f"  {k}: {v}")
    print(f"\nImágenes configuradas: {len(imagenes)}")
    for k, v in imagenes.items():
        estado = "✓" if v.exists() else "✗ NO ENCONTRADA"
        print(f"  {k}: {v.name}  {estado}")
    print()

    # Cargar template y aplicar reemplazos
    doc = Document(str(template_path))

    txt_count = apply_text_replacements(doc, datos)
    img_count, img_errores = apply_image_replacements(doc, imagenes, args.ancho)

    # Guardar
    doc.save(str(salida_path))

    print(f"\n{'─'*50}")
    print(f"✅ Texto reemplazado : {txt_count} placeholder(s)")
    print(f"🖼  Imágenes insertadas: {img_count} de {len(imagenes)}")
    if img_errores:
        print("\nAdvertencias:")
        for e in img_errores:
            print(e)
    print(f"\n📁 Documento generado: {salida_path}\n")


if __name__ == "__main__":
    main()
